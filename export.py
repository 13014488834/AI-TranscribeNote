"""导出层 — Markdown / Word / PDF"""
import os
from datetime import datetime
from config import EXPORT_DIR

# 可选依赖
try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from fpdf import FPDF
    PDF_OK = True
except ImportError:
    PDF_OK = False


def _build_export_content(result: dict) -> str:
    """构建统一的导出文本内容（Markdown 格式）"""
    待办 = result.get("待办事项", [])
    todo_lines = ""
    if 待办:
        for i, item in enumerate(待办, 1):
            task = item.get("任务", item.get("task", ""))
            person = item.get("负责人", item.get("person", ""))
            todo_lines += f"{i}. {task}  — 负责人：{person}\n"
    else:
        todo_lines = "（无）\n"

    content = (
        f"# 会议纪要\n\n"
        f"**生成时间**：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        f"---\n\n"
        f"## 争论焦点\n\n{result.get('争论焦点', '')}\n\n"
        f"---\n\n"
        f"## 最终结论\n\n{result.get('最终结论', '')}\n\n"
        f"---\n\n"
        f"## 待办事项\n\n{todo_lines}\n"
    )
    return content


def export_markdown(result: dict) -> str:
    """导出 Markdown 文件，返回文件路径"""
    content = _build_export_content(result)
    filename = f"会议纪要_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    filepath = EXPORT_DIR / filename
    filepath.write_text(content, encoding="utf-8")
    return str(filepath)


def export_docx(result: dict) -> str:
    """导出 Word 文件，返回文件路径"""
    if not DOCX_OK:
        return ""
    doc = DocxDocument()
    doc.styles["Normal"].font.name = "Microsoft YaHei"

    doc.add_heading("会议纪要", level=0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("争论焦点", level=1)
    doc.add_paragraph(result.get("争论焦点", ""))

    doc.add_heading("最终结论", level=1)
    doc.add_paragraph(result.get("最终结论", ""))

    doc.add_heading("待办事项", level=1)
    待办 = result.get("待办事项", [])
    if 待办:
        for i, item in enumerate(待办, 1):
            task = item.get("任务", item.get("task", ""))
            person = item.get("负责人", item.get("person", ""))
            doc.add_paragraph(f"{i}. {task}  — 负责人：{person}")
    else:
        doc.add_paragraph("（无）")

    filename = f"会议纪要_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = str(EXPORT_DIR / filename)
    doc.save(filepath)
    return filepath


def export_pdf(result: dict) -> str:
    """导出 PDF 文件，返回文件路径"""
    if not PDF_OK:
        return ""

    pdf = FPDF()
    pdf.add_page()

    # 注册中文字体（使用 Windows 系统自带微软雅黑）
    font_path = "C:/Windows/Fonts/msyh.ttc"
    if not os.path.exists(font_path):
        font_path = "C:/Windows/Fonts/msyh.ttf"
    if os.path.exists(font_path):
        pdf.add_font("YaHei", "", font_path)
        pdf.add_font("YaHei", "B", font_path)
    else:
        pdf.add_font("YaHei", "", "")
        pdf.add_font("YaHei", "B", "", "")

    # 标题
    pdf.set_font("YaHei", "B", 18)
    pdf.cell(0, 12, "会议纪要", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)

    # 生成时间
    pdf.set_font("YaHei", "", 9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 8, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
             new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    # 争论焦点
    pdf.set_font("YaHei", "B", 14)
    pdf.cell(0, 10, "争论焦点", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("YaHei", "", 11)
    pdf.multi_cell(0, 7, result.get("争论焦点", ""))
    pdf.ln(4)

    # 最终结论
    pdf.set_font("YaHei", "B", 14)
    pdf.cell(0, 10, "最终结论", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("YaHei", "", 11)
    pdf.multi_cell(0, 7, result.get("最终结论", ""))
    pdf.ln(4)

    # 待办事项
    pdf.set_font("YaHei", "B", 14)
    pdf.cell(0, 10, "待办事项", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("YaHei", "", 11)
    待办 = result.get("待办事项", [])
    if 待办:
        for i, item in enumerate(待办, 1):
            task = item.get("任务", item.get("task", ""))
            person = item.get("负责人", item.get("person", ""))
            pdf.cell(0, 7, f"{i}. {task}  — 负责人：{person}", new_x="LMARGIN", new_y="NEXT")
    else:
        pdf.cell(0, 7, "（无）", new_x="LMARGIN", new_y="NEXT")

    filename = f"会议纪要_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = str(EXPORT_DIR / filename)
    pdf.output(filepath)
    return filepath
