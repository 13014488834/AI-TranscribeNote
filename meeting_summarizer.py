"""
AI 会议纪要结构化生成工具
基于 Gradio + DeepSeek API + LangChain + Pydantic

功能：
  1. 文件上传 — 支持 .txt / .docx / .mp3（Whisper 语音转文字）
  2. 历史记录 — SQLite 本地存储，左侧面板检索与回看
  3. 导出功能 — 一键导出 Markdown / Word / PDF
"""

import subprocess
import sys
import importlib.util

# ==================== 自动安装依赖 ====================
DEP_MAP = {
    "gradio": "gradio",
    "langchain-deepseek": "langchain_deepseek",
    "python-dotenv": "dotenv",
    "pydantic": "pydantic",
    "python-docx": "docx",           # Word 读写
    "fpdf2": "fpdf",                 # PDF 导出
    "openai-whisper": "whisper",     # 语音转文字（可选）
}


def ensure_dependencies() -> None:
    """检查并自动安装缺失的 Python 包"""
    for pip_name, module_name in DEP_MAP.items():
        if importlib.util.find_spec(module_name) is None:
            tag = "（可选）" if pip_name == "openai-whisper" else ""
            print(f"[依赖检查] 缺少 {pip_name} {tag}，正在安装...")
            try:
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", pip_name],
                    stdout=sys.stdout,
                    stderr=sys.stderr,
                )
                print(f"[依赖检查] {pip_name} 安装完成。")
            except Exception as e:
                if pip_name == "openai-whisper":
                    print(f"[依赖检查] openai-whisper 安装失败，音频功能不可用: {e}")
                else:
                    raise
        else:
            print(f"[依赖检查] {pip_name} 已安装，跳过。")


ensure_dependencies()

# ==================== 导入依赖 ====================
import os
import json
import sqlite3
import uuid
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple
from threading import Lock

import gradio as gr
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from langchain_deepseek import ChatDeepSeek

# --- 可选依赖 ---
try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import whisper
    WHISPER_OK = True
except ImportError:
    WHISPER_OK = False

try:
    from fpdf import FPDF
    PDF_OK = True
except ImportError:
    PDF_OK = False

# ==================== 初始化 ====================
load_dotenv()
BASE_DIR = Path(__file__).resolve().parent
DB_PATH = str(BASE_DIR / "meetings.db")
EXPORT_DIR = BASE_DIR / "exports"
EXPORT_DIR.mkdir(exist_ok=True)
DB_LOCK = Lock()

# whisper 模型懒加载（首次使用语音转文字时才下载模型）
_whisper_model = None


def _get_whisper_model():
    """懒加载 Whisper 模型（tiny 模型约 75MB，兼顾速度与准确度）"""
    global _whisper_model
    if _whisper_model is None and WHISPER_OK:
        print("[Whisper] 正在加载语音识别模型 (tiny)...")
        _whisper_model = whisper.load_model("tiny")
        print("[Whisper] 模型加载完成。")
    return _whisper_model


# ==================== Pydantic 数据模型 ====================
class TodoItem(BaseModel):
    """待办事项"""
    task: str = Field(description="任务描述")
    person: str = Field(description="负责人")


class MeetingMinutes(BaseModel):
    """会议纪要结构化模型，LLM 强制按此 schema 返回 JSON"""
    争论焦点: str = Field(description="会议中意见不一致的核心问题")
    最终结论: str = Field(description="会议达成的最终决定")
    待办事项: List[TodoItem] = Field(description="待办事项列表，每项包含任务描述和负责人")


# ==================== LLM 初始化 ====================
_api_key_from_env = os.getenv("DEEPSEEK_API_KEY")
_llm_cache = {}  # 不同 API Key 对应不同 LLM 实例


def _get_structured_llm(override_key: str = ""):
    """懒加载 DeepSeek LLM（首次调用时初始化，支持 UI 临时 Key 或 .env Key）

    Args:
        override_key: 用户在 UI 中输入的 API Key，优先级高于 .env
    """
    effective_key = (override_key or _api_key_from_env or "").strip()

    if not effective_key:
        return None, (
            "⚠️ 未配置 DeepSeek API Key。\n"
            "两种方式任选：\n"
            "1) 在下方「API Key」输入框粘贴你的 Key\n"
            "2) 创建 .env 文件，写入 DEEPSEEK_API_KEY=你的Key\n"
            "获取地址：https://platform.deepseek.com"
        )

    # 检查缓存（同一 Key 复用）
    if effective_key in _llm_cache:
        return _llm_cache[effective_key], None

    try:
        llm = ChatDeepSeek(model="deepseek-chat", temperature=0.1, api_key=effective_key)
        _llm_cache[effective_key] = llm.with_structured_output(MeetingMinutes)
        return _llm_cache[effective_key], None
    except Exception as e:
        return None, f"❌ LLM 初始化失败：{e}"


# ==================== SQLite 数据库层 ====================
def init_db() -> None:
    """创建数据库和表（首次运行自动执行）"""
    with DB_LOCK:
        conn = sqlite3.connect(DB_PATH)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS meetings (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                created_at  TEXT    NOT NULL,
                source_type TEXT    NOT NULL,
                source_name TEXT    NOT NULL,
                original_text TEXT  NOT NULL,
                debate_points   TEXT NOT NULL,
                final_conclusion TEXT NOT NULL,
                todo_items      TEXT NOT NULL
            )
        """)
        conn.commit()
        conn.close()


def save_meeting(
    source_type: str,
    source_name: str,
    original_text: str,
    debate_points: str,
    final_conclusion: str,
    todo_items: list,
) -> int:
    """保存纪要，返回记录 ID"""
    with DB_LOCK:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.execute(
            """INSERT INTO meetings
               (created_at, source_type, source_name, original_text,
                debate_points, final_conclusion, todo_items)
               VALUES (?, ?, ?, ?, ?, ?, ?)""",
            (
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                source_type,
                source_name,
                original_text,
                debate_points,
                final_conclusion,
                json.dumps(todo_items, ensure_ascii=False),
            ),
        )
        row_id = cursor.lastrowid
        conn.commit()
        conn.close()
        return row_id


def get_history(search: str = "") -> list:
    """查询历史记录，返回 [(id, created_at, source_name, debate_points), ...]"""
    with DB_LOCK:
        conn = sqlite3.connect(DB_PATH)
        if search.strip():
            cursor = conn.execute(
                """SELECT id, created_at, source_name, debate_points
                   FROM meetings
                   WHERE original_text   LIKE ?
                      OR debate_points   LIKE ?
                      OR final_conclusion LIKE ?
                   ORDER BY id DESC LIMIT 50""",
                (f"%{search}%", f"%{search}%", f"%{search}%"),
            )
        else:
            cursor = conn.execute(
                """SELECT id, created_at, source_name, debate_points
                   FROM meetings ORDER BY id DESC LIMIT 50"""
            )
        rows = cursor.fetchall()
        conn.close()
        return rows


def get_meeting_by_id(meeting_id: int) -> Optional[dict]:
    """根据 ID 获取完整纪要"""
    with DB_LOCK:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.execute("SELECT * FROM meetings WHERE id = ?", (meeting_id,))
        row = cursor.fetchone()
        conn.close()
    if row is None:
        return None
    return {
        "id": row[0],
        "created_at": row[1],
        "source_type": row[2],
        "source_name": row[3],
        "original_text": row[4],
        "debate_points": row[5],
        "final_conclusion": row[6],
        "todo_items": json.loads(row[7]),
    }


# ==================== 文件解析层 ====================
def parse_txt(file_path: str) -> Tuple[str, Optional[str]]:
    """
    读取纯文本文件
    返回: (文本内容, 错误信息)
    """
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        if not content.strip():
            return "", "文件内容为空。"
        return content, None
    except UnicodeDecodeError:
        # 尝试 GBK 编码（Windows 常见）
        try:
            with open(file_path, "r", encoding="gbk") as f:
                return f.read(), None
        except Exception as e:
            return "", f"文件编码错误：{e}"
    except Exception as e:
        return "", f"读取文件失败：{e}"


def parse_docx(file_path: str) -> Tuple[str, Optional[str]]:
    """
    读取 Word 文档，提取所有段落文本
    返回: (文本内容, 错误信息)
    """
    if not DOCX_OK:
        return "", "缺少 python-docx 依赖，无法解析 .docx 文件。"
    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            return "", "Word 文档中未找到文本内容。"
        return "\n".join(paragraphs), None
    except Exception as e:
        return "", f"解析 Word 文档失败：{e}"


def transcribe_audio(file_path: str) -> Tuple[str, Optional[str]]:
    """
    使用 Whisper 将音频转为文字
    返回: (转写文本, 错误信息)
    """
    if not WHISPER_OK:
        return "", "语音转文字需要 openai-whisper 依赖，安装失败或已跳过。"
    try:
        model = _get_whisper_model()
        if model is None:
            return "", "Whisper 模型加载失败。"
        result = model.transcribe(file_path, fp16=False, language="zh")
        text = result["text"].strip()
        if not text:
            return "", "音频中未检测到语音内容。"
        return text, None
    except Exception as e:
        msg = str(e)
        if "ffmpeg" in msg.lower():
            return "", "需要安装 ffmpeg 才能转写音频。请从 https://ffmpeg.org/download.html 下载并添加到 PATH。"
        return "", f"语音转写失败：{e}"


def parse_file(file_path: str, file_name: str) -> Tuple[str, str, Optional[str]]:
    """
    根据文件扩展名自动选择解析器
    返回: (文本内容, source_name, 错误信息)
    """
    ext = Path(file_name).suffix.lower()
    if ext == ".txt":
        text, err = parse_txt(file_path)
    elif ext == ".docx":
        text, err = parse_docx(file_path)
    elif ext in (".mp3", ".wav", ".m4a", ".flac", ".ogg"):
        text, err = transcribe_audio(file_path)
    else:
        return "", file_name, f"不支持的文件格式「{ext}」。支持的格式：.txt / .docx / .mp3 / .wav / .m4a"
    return text, file_name, err


# ==================== 导出层 ====================
def _build_export_content(result: dict) -> str:
    """构建统一的导出文本内容（Markdown 格式）"""
    待办 = result.get("待办事项", [])
    todo_lines = ""
    if 待办:
        for i, item in enumerate(待办, 1):
            task = item.get("任务", item.get("task", ""))
            person = item.get("负责人", item.get("person", ""))
            todo_lines += f"{i}. {task}  — 负责人：{person}\n"
    else:
        todo_lines = "（无）\n"

    content = (
        f"# 会议纪要\n\n"
        f"**生成时间**：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        f"---\n\n"
        f"## 争论焦点\n\n{result.get('争论焦点', '')}\n\n"
        f"---\n\n"
        f"## 最终结论\n\n{result.get('最终结论', '')}\n\n"
        f"---\n\n"
        f"## 待办事项\n\n{todo_lines}\n"
    )
    return content


def export_markdown(result: dict) -> str:
    """导出 Markdown 文件，返回文件路径"""
    content = _build_export_content(result)
    filename = f"会议纪要_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    filepath = EXPORT_DIR / filename
    filepath.write_text(content, encoding="utf-8")
    return str(filepath)


def export_docx(result: dict) -> str:
    """导出 Word 文件，返回文件路径"""
    if not DOCX_OK:
        return ""
    doc = DocxDocument()
    doc.styles["Normal"].font.name = "Microsoft YaHei"

    doc.add_heading("会议纪要", level=0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("争论焦点", level=1)
    doc.add_paragraph(result.get("争论焦点", ""))

    doc.add_heading("最终结论", level=1)
    doc.add_paragraph(result.get("最终结论", ""))

    doc.add_heading("待办事项", level=1)
    待办 = result.get("待办事项", [])
    if 待办:
        for i, item in enumerate(待办, 1):
            task = item.get("任务", item.get("task", ""))
            person = item.get("负责人", item.get("person", ""))
            doc.add_paragraph(f"{i}. {task}  — 负责人：{person}")
    else:
        doc.add_paragraph("（无）")

    filename = f"会议纪要_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = str(EXPORT_DIR / filename)
    doc.save(filepath)
    return filepath


def export_pdf(result: dict) -> str:
    """导出 PDF 文件，返回文件路径"""
    if not PDF_OK:
        return ""

    pdf = FPDF()
    pdf.add_page()

    # 注册中文字体（使用 Windows 系统自带微软雅黑）
    font_path = "C:/Windows/Fonts/msyh.ttc"
    if not os.path.exists(font_path):
        font_path = "C:/Windows/Fonts/msyh.ttf"
    if os.path.exists(font_path):
        pdf.add_font("YaHei", "", font_path, uni=True)
        pdf.add_font("YaHei", "B", font_path, uni=True)
    else:
        # 降级：无中文字体时用内置字体（中文会显示为方块）
        pdf.add_font("YaHei", "", "")
        pdf.add_font("YaHei", "B", "", "")

    # 标题
    pdf.set_font("YaHei", "B", 18)
    pdf.cell(0, 12, "会议纪要", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)

    # 生成时间
    pdf.set_font("YaHei", "", 9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 8, f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
             new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    # 争论焦点
    pdf.set_font("YaHei", "B", 14)
    pdf.cell(0, 10, "争论焦点", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("YaHei", "", 11)
    pdf.multi_cell(0, 7, result.get("争论焦点", ""))
    pdf.ln(4)

    # 最终结论
    pdf.set_font("YaHei", "B", 14)
    pdf.cell(0, 10, "最终结论", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("YaHei", "", 11)
    pdf.multi_cell(0, 7, result.get("最终结论", ""))
    pdf.ln(4)

    # 待办事项
    pdf.set_font("YaHei", "B", 14)
    pdf.cell(0, 10, "待办事项", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("YaHei", "", 11)
    待办 = result.get("待办事项", [])
    if 待办:
        for i, item in enumerate(待办, 1):
            task = item.get("任务", item.get("task", ""))
            person = item.get("负责人", item.get("person", ""))
            pdf.cell(0, 7, f"{i}. {task}  — 负责人：{person}", new_x="LMARGIN", new_y="NEXT")
    else:
        pdf.cell(0, 7, "（无）", new_x="LMARGIN", new_y="NEXT")

    filename = f"会议纪要_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = str(EXPORT_DIR / filename)
    pdf.output(filepath)
    return filepath


# ==================== 核心处理函数 ====================
def summarize_meeting(text: str, source_type: str = "text",
                      source_name: str = "手动输入",
                      api_key: str = "") -> dict:
    """
    处理会议文字记录，提取争论焦点、最终结论和待办事项，
    并自动保存到 SQLite 历史记录。

    Args:
        text: 会议文字内容
        source_type: 来源类型（text / file）
        source_name: 来源名称（用于历史列表展示）
        api_key: 用户在 UI 输入的 DeepSeek API Key，优先级高于 .env

    Returns:
        dict: 结构化纪要
    """
    # 空输入检查
    if not text or not text.strip():
        return {
            "状态": "⚠️ 输入为空",
            "提示": "请粘贴会议内容或上传文件。",
            "争论焦点": "",
            "最终结论": "",
            "待办事项": [],
        }

    # 文本过长时截断提示
    MAX_CHARS = 8000
    overflow_note = ""
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS]
        overflow_note = "「注意：原文过长，已截取前8000字处理」"

    # 构建提示词
    prompt = (
        "你是一位资深的会议纪要秘书。请仔细阅读以下会议文字记录，从中提取出：\n"
        "1. 争论焦点：会议中意见不一致、存在讨论或争议的核心问题是什么？\n"
        "2. 最终结论：会议针对上述争论最终达成了什么决定或共识？\n"
        "3. 待办事项：列出所有需要后续跟进的任务，每项需明确任务描述和负责人。\n\n"
        "重要规则：\n"
        "- 如果某项信息在会议记录中没有体现，请填写「（会议中未提及）」，不要凭空编造。\n"
        "- 待办事项如果会议中没提到任何人名，负责人请填写「待定」。\n"
        "- 尽可能精确提取，保持原文关键信息不丢失。\n\n"
        f"{overflow_note}\n=== 会议记录 ===\n{text}"
    )

    # 懒加载 LLM（首次调用时初始化）
    llm_instance, llm_error = _get_structured_llm(override_key=api_key)
    if llm_error:
        return {
            "状态": "⚠️ 未配置 API Key",
            "提示": llm_error,
            "争论焦点": "", "最终结论": "", "待办事项": [],
        }
    if llm_instance is None:
        return {
            "状态": "❌ LLM 未初始化",
            "提示": "内部错误：LLM 初始化失败。",
            "争论焦点": "", "最终结论": "", "待办事项": [],
        }

    # 调用 API
    try:
        result: MeetingMinutes = llm_instance.invoke(prompt)
        todo_list = [
            {"任务": item.task, "负责人": item.person}
            for item in result.待办事项
        ]

        output = {
            "状态": "✅ 纪要生成成功",
            "争论焦点": result.争论焦点,
            "最终结论": result.最终结论,
            "待办事项": todo_list,
        }

        # 自动保存到历史记录
        try:
            save_meeting(
                source_type=source_type,
                source_name=source_name,
                original_text=text,
                debate_points=result.争论焦点,
                final_conclusion=result.最终结论,
                todo_items=todo_list,
            )
        except Exception:
            pass  # 数据库写入失败不影响主流程

        return output

    except Exception as e:
        return {
            "状态": "❌ API 调用失败",
            "错误信息": str(e),
            "争论焦点": "",
            "最终结论": "",
            "待办事项": [],
        }


# ==================== Gradio 事件处理 ====================
def _refresh_history_choices(search: str = "") -> list:
    """刷新历史记录 Radio 的可选项"""
    records = get_history(search)
    if not records:
        return []
    # 格式: "2025-06-26 14:30 | 争论焦点前30字..."
    return [
        (f"{r[1]} | {r[3][:35]}{'...' if len(r[3]) > 35 else ''}", r[0])
        for r in records
    ]


def on_text_submit(text: str, api_key: str = "") -> Tuple[dict, dict, dict]:
    """文本输入 → 生成纪要 + 刷新历史 + 更新导出状态"""
    result = summarize_meeting(text, "text", "手动输入", api_key=api_key)
    choices = _refresh_history_choices()
    return result, gr.update(choices=choices, value=None), result


def on_file_submit(file, api_key: str = "") -> Tuple[dict, dict, dict]:
    """文件上传 → 解析 → 生成纪要 + 刷新历史 + 更新导出状态"""
    empty = {
        "状态": "⚠️ 未选择文件",
        "争论焦点": "", "最终结论": "", "待办事项": [],
    }
    if file is None:
        return empty, gr.update(choices=_refresh_history_choices(), value=None), empty

    file_path = file if isinstance(file, str) else file.name
    file_name = os.path.basename(file_path)

    # 解析文件
    text, name, err = parse_file(file_path, file_name)
    if err:
        empty["状态"] = "❌ 文件解析失败"
        empty["错误信息"] = err
        return empty, gr.update(choices=_refresh_history_choices(), value=None), empty

    # 生成纪要
    result = summarize_meeting(text, "file", name, api_key=api_key)
    choices = _refresh_history_choices()
    return result, gr.update(choices=choices, value=None), result


def on_history_select(selected_str: str) -> Tuple[dict, Optional[str], Optional[str], dict]:
    """用户点击历史记录 → 加载详情"""
    empty: dict = {}
    if not selected_str:
        return empty, None, None, empty

    meeting_id = selected_str

    record = get_meeting_by_id(meeting_id)
    if record is None:
        return {
            "状态": "⚠️ 记录未找到",
            "争论焦点": "", "最终结论": "", "待办事项": [],
        }, None, None, empty

    result = {
        "状态": f"📋 历史记录 #{record['id']} — {record['created_at']}",
        "来源": record["source_name"],
        "争论焦点": record["debate_points"],
        "最终结论": record["final_conclusion"],
        "待办事项": record["todo_items"],
    }
    # DownloadButton 返回 None 表示不触发下载；第4个值更新导出状态
    return result, None, None, result


def on_history_search(query: str) -> dict:
    """搜索历史记录"""
    choices = _refresh_history_choices(query)
    return gr.update(choices=choices, value=None)


def on_export_md(state: dict) -> Optional[str]:
    """导出 Markdown"""
    if not state or not state.get("争论焦点"):
        return None
    return export_markdown(state)


def on_export_docx(state: dict) -> Optional[str]:
    """导出 Word"""
    if not state or not state.get("争论焦点"):
        return None
    return export_docx(state)


def on_export_pdf(state: dict) -> Optional[str]:
    """导出 PDF"""
    if not state or not state.get("争论焦点"):
        return None
    return export_pdf(state)


# ==================== Gradio Web 界面 ====================
# 确保数据库表在 UI 初始化前创建（UI 构建时会查询历史记录）
init_db()

with gr.Blocks(title="AI智能会议纪要生成工具") as demo:
    # 共享状态：保存当前展示的纪要结果，供导出按钮使用
    current_result = gr.State({})

    # --- 标题 ---
    gr.Markdown("""
    # 📋 AI智能会议纪要生成工具
    输入会议文字记录或上传文件，自动提取争论焦点、结论和待办事项
    """)

    # --- API Key 区域（优先 UI 输入，其次 .env） ---
    gr.Markdown("### 🔑 DeepSeek API Key")
    with gr.Row():
        api_key_input = gr.Textbox(
            label="",
            placeholder="粘贴你的 DeepSeek API Key（sk-...），或留空使用 .env 配置",
            type="password",
            scale=3,
            value=_api_key_from_env or "",
        )
        api_key_link = gr.Markdown(
            "[📎 获取 API Key](https://platform.deepseek.com) ｜ 你的 Key 仅在浏览器本地使用，不存储到服务器",
            scale=2,
        )

    with gr.Row():
        # ========== 左侧：历史记录面板 ==========
        with gr.Column(scale=1, min_width=220):
            gr.Markdown("### 📋 历史记录")
            search_box = gr.Textbox(
                label="搜索",
                placeholder="输入关键词筛选……",
                scale=1,
            )
            history_radio = gr.Radio(
                choices=_refresh_history_choices(),
                label="最近记录",
                interactive=True,
            )

        # ========== 右侧：主操作区 ==========
        with gr.Column(scale=3, min_width=500):
            with gr.Tabs():
                # 文本输入 Tab
                with gr.Tab("📝 文本输入"):
                    text_input = gr.Textbox(
                        label="会议文字记录",
                        placeholder="请粘贴会议录音转文字……",
                        lines=12,
                    )
                    text_submit_btn = gr.Button("🚀 生成纪要", variant="primary", size="lg")

                # 文件上传 Tab
                with gr.Tab("📁 文件上传"):
                    file_input = gr.File(
                        label="上传会议文件",
                        file_types=[".txt", ".docx", ".mp3", ".wav", ".m4a", ".flac", ".ogg"],
                    )
                    file_submit_btn = gr.Button("🚀 解析并生成纪要", variant="primary", size="lg")

            # 输出区域
            output_json = gr.JSON(
                label="结构化会议纪要",
            )

            # 导出按钮
            with gr.Row():
                export_md_btn = gr.DownloadButton("📥 导出 Markdown")
                export_docx_btn = gr.DownloadButton("📥 导出 Word")
                export_pdf_btn = gr.DownloadButton("📥 导出 PDF")

    # ==================== 事件绑定 ====================
    # 文本提交 → 同时更新 JSON 显示、历史列表、导出状态
    text_submit_btn.click(
        fn=lambda txt, key: on_text_submit(txt, str(key or "")),
        inputs=[text_input, api_key_input],
        outputs=[output_json, history_radio, current_result],
    )

    # 文件提交 → 同时更新 JSON 显示、历史列表、导出状态
    file_submit_btn.click(
        fn=lambda f, key: on_file_submit(f, str(key or "")),
        inputs=[file_input, api_key_input],
        outputs=[output_json, history_radio, current_result],
    )

    # 历史搜索
    search_box.change(
        fn=on_history_search,
        inputs=search_box,
        outputs=history_radio,
    )

    # 历史点击 → 加载详情 + 更新导出状态
    history_radio.change(
        fn=on_history_select,
        inputs=history_radio,
        outputs=[output_json, export_md_btn, export_docx_btn, current_result],
    )

    # 导出（从 current_result State 取数据）
    export_md_btn.click(
        fn=on_export_md,
        inputs=current_result,
        outputs=export_md_btn,
    )
    export_docx_btn.click(
        fn=on_export_docx,
        inputs=current_result,
        outputs=export_docx_btn,
    )
    export_pdf_btn.click(
        fn=on_export_pdf,
        inputs=current_result,
        outputs=export_pdf_btn,
    )


# ==================== 启动入口 ====================
if __name__ == "__main__":
    init_db()
    print(f"[初始化] 数据库路径：{DB_PATH}")
    print(f"[初始化] 导出目录：{EXPORT_DIR}")

    demo.launch(
        inbrowser=True,
        theme=gr.themes.Soft(),
        css="""
            /* Windows 中文字体渲染 */
            * { font-family: "Microsoft YaHei", "微软雅黑", sans-serif; }
            /* 左侧历史面板微调 */
            .history-panel { max-height: 500px; overflow-y: auto; }
        """,
    )
